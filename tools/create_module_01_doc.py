from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Business-Module-Documents/01-identity-access-management-business-analysis.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.margin_top = 80
            cell.margin_bottom = 80
            cell.margin_left = 120
            cell.margin_right = 120
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
            if row_idx == 0:
                set_cell_shading(cell, "D9EAF7")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(31, 78, 121) if level > 1 else RGBColor(46, 116, 181)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Field", True)
    set_cell_text(hdr[1], "Value", True)
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, True)
        set_cell_text(cells[1], value)
    style_table(table)
    return table


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(table)
    return table


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Module 01: Identity & Access Management")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Business Analysis Document - Example")
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph()
    add_kv_table(
        doc,
        [
            ("Project", "E-Learning Product Owner and System Design Learning Project"),
            ("Module", "Identity & Access Management"),
            ("Document Type", "Business Analysis Document"),
            ("Version", "1.0"),
            ("Audience", "Product Owner, Business Analyst, UI/UX, QA, Software Team"),
        ],
    )

    add_heading(doc, "1. Module Purpose", 1)
    doc.add_paragraph(
        "Identity & Access Management handles user registration, login, roles, permissions, and account access for the e-learning platform."
    )

    add_heading(doc, "2. Scope", 1)
    add_table(
        doc,
        ["In Scope", "Out of Scope"],
        [
            ("Student registration and login", "Social login"),
            ("Parent registration and login", "Two-factor authentication"),
            ("Teacher account approval by admin", "Single sign-on"),
            ("Admin login and user management", "Advanced identity provider integration"),
            ("Role-based access control", "Biometric authentication"),
            ("Basic profile update", "Complex account recovery workflow"),
        ],
    )

    add_heading(doc, "3. Actors", 1)
    add_table(
        doc,
        ["Actor", "Need"],
        [
            ("Student", "Create account, log in, access enrolled learning content"),
            ("Parent", "Create account, link to student, view progress and redeem codes"),
            ("Teacher", "Access teacher dashboard after approval"),
            ("Admin", "Approve teachers and manage user accounts"),
        ],
    )

    add_heading(doc, "4. Business Capabilities", 1)
    add_bullets(
        doc,
        [
            "Register student and parent accounts.",
            "Allow users to log in and log out.",
            "Support user roles: Student, Parent, Teacher, Admin.",
            "Allow admin to approve or reject teacher accounts.",
            "Control access based on role and permissions.",
            "Allow users to update basic profile information.",
        ],
    )

    add_heading(doc, "5. Business Rules", 1)
    add_table(
        doc,
        ["Rule ID", "Business Rule"],
        [
            ("IAM-BR-01", "Students can register without parent approval."),
            ("IAM-BR-02", "Parents can register and link to students using student ID."),
            ("IAM-BR-03", "Teachers cannot publish courses before admin approval."),
            ("IAM-BR-04", "Each user must have one primary role in MVP."),
            ("IAM-BR-05", "Users can access only screens allowed for their role."),
            ("IAM-BR-06", "Admin actions should be recorded for audit."),
        ],
    )

    add_heading(doc, "6. Functional Requirements", 1)
    add_table(
        doc,
        ["ID", "Requirement"],
        [
            ("IAM-FR-01", "The system shall allow students to create accounts."),
            ("IAM-FR-02", "The system shall allow parents to create accounts."),
            ("IAM-FR-03", "The system shall allow users to log in and log out."),
            ("IAM-FR-04", "The system shall allow users to update basic profile data."),
            ("IAM-FR-05", "The system shall allow teacher accounts to be submitted for approval."),
            ("IAM-FR-06", "The system shall allow admin to approve or reject teachers."),
            ("IAM-FR-07", "The system shall enforce role-based access to screens and actions."),
            ("IAM-FR-08", "The system shall prevent duplicate registration using the same phone or email."),
        ],
    )

    add_heading(doc, "7. User Stories", 1)
    add_table(
        doc,
        ["Story ID", "User Story", "Acceptance Criteria"],
        [
            (
                "IAM-US-01",
                "As a student, I want to register so that I can use the platform.",
                "Student enters required data; duplicate phone/email is rejected; account is created.",
            ),
            (
                "IAM-US-02",
                "As a parent, I want to register so that I can follow my child.",
                "Parent enters required data; account is created; parent can log in.",
            ),
            (
                "IAM-US-03",
                "As a teacher, I want my account approved so that I can create courses.",
                "Teacher submits profile data; admin approves/rejects; rejected teacher cannot publish.",
            ),
            (
                "IAM-US-04",
                "As an admin, I want to manage user accounts so that platform access is controlled.",
                "Admin can view users; update status; approve teachers.",
            ),
        ],
    )

    add_heading(doc, "8. Screens", 1)
    add_table(
        doc,
        ["Screen", "User", "Purpose"],
        [
            ("Register", "Student / Parent", "Create new account"),
            ("Login", "All users", "Access platform account"),
            ("Profile", "All users", "View and update basic profile data"),
            ("Teacher Approval", "Admin", "Review and approve/reject teachers"),
            ("User Management", "Admin", "Manage account status and user list"),
        ],
    )

    add_heading(doc, "9. Data Objects", 1)
    add_table(
        doc,
        ["Object", "Important Fields"],
        [
            ("User", "ID, name, phone, email, password, role, status"),
            ("Student Profile", "User ID, student ID, secondary year"),
            ("Parent Profile", "User ID"),
            ("Teacher Profile", "User ID, subject, bio, profile image, documents, approval status"),
            ("Audit Log", "Action, user ID, target ID, date/time"),
        ],
    )

    add_heading(doc, "10. Test Scenarios", 1)
    add_table(
        doc,
        ["ID", "Scenario", "Expected Result"],
        [
            ("IAM-TS-01", "Student registers with valid data", "Student account is created."),
            ("IAM-TS-02", "User logs in with valid credentials", "User enters correct dashboard."),
            ("IAM-TS-03", "Student registers with existing phone/email", "System rejects registration."),
            ("IAM-TS-04", "Admin approves teacher", "Teacher can access teacher features."),
            ("IAM-TS-05", "Rejected teacher logs in", "Teacher cannot publish courses."),
            ("IAM-TS-06", "User opens unauthorized screen", "System blocks access."),
        ],
    )

    add_heading(doc, "11. Open Questions", 1)
    add_bullets(
        doc,
        [
            "Is phone number required for all users?",
            "Should email be optional for students?",
            "Should teacher approval require document upload in MVP?",
            "Should users be deactivated or deleted when removed?",
        ],
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("E-Learning Platform - Module Business Document Example")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(128, 128, 128)

    doc.save(OUT)


if __name__ == "__main__":
    main()

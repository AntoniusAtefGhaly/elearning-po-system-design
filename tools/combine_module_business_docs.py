from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MODULES_DIR = ROOT / "Modules"
OUTPUT = ROOT / "Modules" / "E-Learning Platform - Combined Module Business Document.docx"


def iter_block_items(doc):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield "paragraph", child
        elif child.tag == qn("w:tbl"):
            yield "table", child


def paragraph_text(paragraph_element):
    texts = []
    for node in paragraph_element.iter(qn("w:t")):
        if node.text:
            texts.append(node.text)
    return "".join(texts).strip()


def table_rows(table_element):
    rows = []
    for tr in table_element.findall(qn("w:tr")):
        cells = []
        for tc in tr.findall(qn("w:tc")):
            cell_text = []
            for t in tc.iter(qn("w:t")):
                if t.text:
                    cell_text.append(t.text)
            cells.append(" ".join(cell_text).strip())
        if any(cells):
            rows.append(cells)
    return rows


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)


def add_source_table(out_doc, rows):
    if not rows:
        return

    max_cols = min(max(len(row) for row in rows), 4)
    normalized = []
    for row in rows:
        normalized.append((row + [""] * max_cols)[:max_cols])

    table = out_doc.add_table(rows=len(normalized), cols=max_cols)
    table.style = "Table Grid"
    table.autofit = True
    for row_index, row in enumerate(normalized):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            set_cell_text(cell, value, bold=row_index == 0)
            if row_index == 0:
                set_cell_shading(cell, "D9EAF7")


def module_title(path):
    name = path.stem
    cleanup = {
        "00 E-Learning Platform - Master Business Document (2) (2)": "00 - Master Business Document",
        "01-identity-access-management": "01 - Identity & Access Management",
        "02-TM_BS_rewritten_same_design": "02 - Teacher Management",
        "03-Student_Management_module": "03 - Student Management",
        "04_Content_Course_Managemen": "04 - Content & Course Management",
        "05-Enrollment Management": "05 - Enrollment Management",
        "06 -Assessment Management": "06 - Assessment Management",
        "07_Grading_Management": "07 - Grading Management",
    }
    return cleanup.get(name, name.replace("_", " "))


def configure_styles(doc):
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    for style_name, size, color in [
        ("Title", 24, "1F4E79"),
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 13, "2F5597"),
        ("Heading 3", 11, "1F4E79"),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def add_cover(doc, module_files):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("E-Learning Platform")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Combined Module Business Document")
    subtitle_run.bold = True
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(47, 85, 151)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run("Developer reference compiled from the module business documents.").italic = True

    doc.add_paragraph()
    doc.add_heading("Included Modules", level=1)
    for path in module_files:
        doc.add_paragraph(module_title(path), style="List Bullet")

    doc.add_paragraph()
    doc.add_paragraph(
        "Note: This document combines source module content for developer reference. "
        "If a module contains future-scope items, implementation priority should still follow the product backlog and sprint plan."
    )


def add_module(doc, path, first_module=False):
    if not first_module:
        doc.add_section(WD_SECTION.NEW_PAGE)

    doc.add_heading(module_title(path), level=1)
    doc.add_paragraph(f"Source file: {path.name}")

    source = Document(path)
    for block_type, element in iter_block_items(source):
        if block_type == "paragraph":
            text = paragraph_text(element)
            if not text:
                continue
            if len(text) <= 90 and (
                text.endswith("Document")
                or text.startswith("Module")
                or text[:2].isdigit()
                or text in {
                    "Purpose",
                    "Scope",
                    "Actors",
                    "Business Rules",
                    "Functional Requirements",
                    "User Stories",
                    "Data Objects",
                    "Workflow",
                    "Future Enhancements",
                }
            ):
                doc.add_heading(text, level=2)
            else:
                doc.add_paragraph(text)
        else:
            add_source_table(doc, table_rows(element))


def main():
    module_files = sorted(MODULES_DIR.glob("*.docx"))
    module_files = [p for p in module_files if p.name != OUTPUT.name and not p.name.startswith("~$")]
    if not module_files:
        raise SystemExit("No module docx files found.")

    doc = Document()
    configure_styles(doc)

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    add_cover(doc, module_files)
    for index, path in enumerate(module_files):
        add_module(doc, path, first_module=index == 0)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
